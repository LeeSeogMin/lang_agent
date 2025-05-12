"""
다양한 형식의 문서를 파싱하는 기능
"""
from typing import List, Generator
import os
import fitz  # PyMuPDF
from docx import Document
import logging
from ..config.settings import DOCUMENTS_DIR
import pytesseract
from pdf2image import convert_from_path
import tempfile
from pathlib import Path

# 로깅 설정
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

class DocumentParser:
    def parse(self, file_path: str, chunk_size: int = 1000) -> List[str]:
        """파일 형식에 따라 적절한 파서를 선택하여 문서를 파싱"""
        ext = os.path.splitext(file_path)[1].lower()
        
        logger.info(f"문서 파싱 시작: {file_path} (크기: {os.path.getsize(file_path)} bytes)")
        
        if ext == '.pdf':
            return list(self.parse_pdf(file_path, chunk_size))
        elif ext == '.txt':
            return list(self.parse_txt(file_path, chunk_size))
        elif ext == '.docx':
            return list(self.parse_docx(file_path, chunk_size))
        else:
            raise ValueError(f"지원하지 않는 파일 형식: {ext}")
    
    def parse_pdf(self, file_path: str, chunk_size: int) -> Generator[str, None, None]:
        """PDF 파일 파싱"""
        try:
            logger.info(f"PDF 파싱 시작: {file_path}")
            doc = fitz.open(file_path)
            logger.info(f"PDF 문서 열기 성공: {len(doc)} 페이지")
            
            total_text = ""
            needs_ocr = True  # OCR 필요 여부 플래그
            
            # 먼저 PyMuPDF로 텍스트 추출 시도
            for page_num, page in enumerate(doc):
                try:
                    logger.debug(f"페이지 {page_num+1} 처리 중...")
                    text = page.get_text()
                    if text.strip():
                        needs_ocr = False  # 텍스트가 추출되면 OCR 불필요
                        total_text += f"\n[페이지 {page_num+1}]\n{text.strip()}"
                except Exception as page_error:
                    logger.error(f"페이지 {page_num+1} 처리 중 오류: {str(page_error)}")
            
            doc.close()
            
            # PyMuPDF로 텍스트 추출 실패 시 OCR 시도
            if needs_ocr:
                logger.info("텍스트 추출 실패, OCR 시작...")
                try:
                    # PDF를 이미지로 변환
                    with tempfile.TemporaryDirectory() as temp_dir:
                        images = convert_from_path(file_path)
                        for i, image in enumerate(images):
                            logger.debug(f"OCR 처리 중: 페이지 {i+1}")
                            # OCR 수행
                            text = pytesseract.image_to_string(image, lang='kor+eng')
                            if text.strip():
                                total_text += f"\n[페이지 {i+1}]\n{text.strip()}"
                except Exception as ocr_error:
                    logger.error(f"OCR 처리 중 오류: {str(ocr_error)}")
            
            logger.info("PDF 문서 파싱 완료, 청킹 시작")
            
            if not total_text.strip():
                logger.warning("추출된 텍스트가 없습니다")
                return
            
            # 전체 텍스트를 청크로 분할
            chunks = list(self.chunk_text(total_text, chunk_size))
            logger.info(f"청킹 완료: {len(chunks)}개 청크 생성")
            
            for chunk in chunks:
                yield chunk
                
        except Exception as e:
            logger.error(f"PDF 파싱 오류 ({file_path}): {str(e)}", exc_info=True)
            raise ValueError(f"PDF 파일 파싱 중 오류 발생: {str(e)}")
    
    def parse_txt(self, file_path: str, chunk_size: int) -> Generator[str, None, None]:
        """텍스트 파일 파싱"""
        try:
            logger.info(f"텍스트 파일 파싱 시작: {file_path}")
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
                if text.strip():
                    line_count = 0
                    chunks = list(self.chunk_text(text, chunk_size))
                    logger.info(f"텍스트 파일 청킹 완료: {len(chunks)}개 청크")
                    for chunk in chunks:
                        yield f"[텍스트 {line_count+1}] {chunk}"
                        line_count += chunk.count('\n') + 1
        except UnicodeDecodeError:
            logger.warning("UTF-8 디코딩 실패, CP949 시도")
            try:
                with open(file_path, 'r', encoding='cp949') as file:
                    text = file.read()
                    if text.strip():
                        chunks = list(self.chunk_text(text, chunk_size))
                        logger.info(f"CP949 인코딩으로 파싱 성공: {len(chunks)}개 청크")
                        for chunk in chunks:
                            yield chunk
            except Exception as e:
                logger.error(f"텍스트 파일 인코딩 오류 ({file_path}): {str(e)}")
                raise ValueError(f"텍스트 파일 파싱 중 오류 발생: 인코딩 문제")
        except Exception as e:
            logger.error(f"텍스트 파일 파싱 오류 ({file_path}): {str(e)}")
            raise ValueError(f"텍스트 파일 파싱 중 오류 발생: {str(e)}")
    
    def parse_docx(self, file_path: str, chunk_size: int) -> Generator[str, None, None]:
        """DOCX 파일 파싱"""
        try:
            logger.info(f"DOCX 파일 파싱 시작: {file_path}")
            doc = Document(file_path)
            sections = []
            
            # 단락 기준으로 처리
            current_section = []
            current_section_len = 0
            
            for para in doc.paragraphs:
                if para.text.strip():
                    # 단락 추가
                    current_section.append(para.text.strip())
                    current_section_len += len(para.text)
                    
                    # 섹션 크기가 청크 크기에 도달하면 처리
                    if current_section_len >= chunk_size:
                        sections.append(' '.join(current_section))
                        current_section = []
                        current_section_len = 0
            
            # 남은 섹션 처리
            if current_section:
                sections.append(' '.join(current_section))
            
            logger.info(f"DOCX 파싱 완료: {len(sections)}개 섹션")
            
            # 섹션 반환
            for i, section in enumerate(sections):
                yield f"[섹션 {i+1}] {section}"
                
            # 섹션이 없으면 전체 텍스트를 하나의 청크로 처리
            if not sections:
                full_text = ' '.join([para.text.strip() for para in doc.paragraphs if para.text.strip()])
                if full_text:
                    chunks = list(self.chunk_text(full_text, chunk_size))
                    logger.info(f"전체 텍스트 청킹: {len(chunks)}개 청크")
                    for chunk in chunks:
                        yield chunk
                        
        except Exception as e:
            logger.error(f"DOCX 파일 파싱 오류 ({file_path}): {str(e)}")
            raise ValueError(f"DOCX 파일 파싱 중 오류 발생: {str(e)}")
    
    def chunk_text(
        self,
        text: str,
        chunk_size: int = 1000,
        overlap: int = 200
    ) -> Generator[str, None, None]:
        """텍스트를 청크로 분할"""
        if not text.strip():
            logger.warning("청킹할 텍스트가 비어있습니다")
            return
        
        logger.debug(f"텍스트 청킹 시작: 총 {len(text)} 문자")
        start = 0
        text_len = len(text)
        chunk_count = 0
        
        while start < text_len:
            end = min(start + chunk_size, text_len)
            
            # 문장 경계에서 끊기
            if end < text_len:
                # 다음 마침표 위치 찾기
                next_period = text.find('.', end - overlap)
                if next_period != -1 and next_period < end + overlap:
                    end = next_period + 1
            
            chunk = text[start:end].strip()
            if chunk:
                chunk_count += 1
                yield chunk
            
            if end >= text_len:
                break
                
            start = end - overlap
        
        logger.debug(f"청킹 완료: {chunk_count}개 청크 생성") 